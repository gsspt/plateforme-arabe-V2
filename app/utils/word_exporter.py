# app/utils/word_exporter.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import logging
from datetime import datetime
from ..config import Config

logger = logging.getLogger(__name__)

class WordExporter:
    """Exportateur de documents Word avec mise en page préservée"""
    
    @staticmethod
    def export_ocr_to_word(ocr_text, filename):
        """Exporte le texte OCR en Word avec mise en page préservée"""
        try:
            logger.info(f"📝 Début export OCR Word: {filename}")
            logger.info(f"📄 Texte à exporter: {len(ocr_text)} caractères")
            
            # Créer le document
            doc = Document()
            
            # Configuration de la page
            section = doc.sections[0]
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # Style par défaut
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # En-tête du document
            title = doc.add_heading('Extraction OCR - Texte Arabe', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(0, 0, 128)
            title_run.font.size = Pt(16)
            
            # Métadonnées
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            meta_run = meta_para.add_run(f"Date d'export: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph().add_run('=' * 60).bold = True
            
            # ✅ TRAITER CHAQUE LIGNE INDIVIDUELLEMENT
            lines = ocr_text.split('\n')
            current_page = None
            
            logger.info(f"📊 Traitement de {len(lines)} lignes...")
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                if not line:
                    # Ligne vide = saut de ligne
                    doc.add_paragraph()
                    continue
                
                # Détecter les en-têtes de page
                if line.startswith('--- Page') and '---' in line:
                    current_page = line.replace('---', '').strip()
                    
                    # Saut de page (sauf pour la première page)
                    if i > 0:
                        doc.add_page_break()
                    
                    # En-tête de page stylisé
                    page_header = doc.add_paragraph()
                    page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    page_header.paragraph_format.space_before = Pt(12)
                    page_header.paragraph_format.space_after = Pt(12)
                    
                    page_run = page_header.add_run(f"📄 {current_page}")
                    page_run.bold = True
                    page_run.font.size = Pt(14)
                    page_run.font.color.rgb = RGBColor(0, 0, 128)
                    
                    # Ligne de séparation
                    sep_para = doc.add_paragraph()
                    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sep_run = sep_para.add_run('―' * 40)
                    sep_run.font.color.rgb = RGBColor(200, 200, 200)
                    sep_run.font.size = Pt(8)
                    
                else:
                    # Texte normal en arabe
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # RTL pour arabe
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.paragraph_format.space_before = Pt(0)
                    
                    # Configuration RTL pour le paragraphe
                    WordExporter._set_paragraph_rtl(paragraph)
                    
                    run = paragraph.add_run(line)
                    run.font.name = 'Traditional Arabic'  # Police optimisée arabe
                    run.font.size = Pt(11)
                    
                    # Si la ligne semble être un titre/sous-titre
                    if (len(line) < 100 and 
                        (line.endswith(':') or not any(c.isalpha() for c in line))):
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 100, 0)
            
            # Pied de page
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("Exporté par Plateforme Éducative Arabe V2")
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            footer_run.italic = True
            
            # ✅ SAUVEGARDE AVEC GESTION DES CHEMINS ABSOLUS
            filepath = os.path.join(Config.UPLOAD_FOLDER, filename)
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            doc.save(filepath)
            logger.info(f"✅ Export OCR Word réussi: {filepath}")
            logger.info(f"📁 Taille du fichier: {os.path.getsize(filepath)} bytes")
            
            return filepath
            
        except Exception as e:
            logger.error(f"❌ Erreur export OCR Word: {e}", exc_info=True)
            raise

    @staticmethod
    def export_translation_to_word(original_text, translated_text, context, filename):
        """Exporte la traduction en Word avec mise en page bilingue"""
        try:
            logger.info(f"🌍 Début export traduction Word: {filename}")
            logger.info(f"📊 Stats: original={len(original_text)}, traduit={len(translated_text)}")
            
            doc = Document()
            
            # Configuration de la page
            section = doc.sections[0]
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # En-tête du document
            title = doc.add_heading('Traduction Arabe → Français', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(0, 100, 0)
            title_run.font.size = Pt(18)
            
            # Métadonnées contextuelles
            meta_table = doc.add_table(rows=1, cols=2)
            meta_table.style = 'Light Grid Accent 1'
            meta_table.autofit = True
            
            meta_cells = meta_table.rows[0].cells
            meta_cells[0].text = "Contexte de traduction:"
            meta_cells[1].text = (f"Auteur: {context.get('auteur', 'Non spécifié')}\n"
                                 f"Titre: {context.get('titre', 'Non spécifié')}\n"
                                 f"Sujet: {context.get('sujet', 'Non spécifié')}\n"
                                 f"Genre: {context.get('genre', 'Non spécifié')}\n"
                                 f"Niveau: {context.get('niveau_langue', 'Standard')}")
            
            doc.add_paragraph().add_run('=' * 60).bold = True
            
            # ✅ TEXTE ORIGINAL ARABE (RTL)
            doc.add_heading('📜 Texte Original (Arabe)', level=1)
            WordExporter._add_formatted_text(doc, original_text, is_arabic=True)
            
            doc.add_page_break()
            
            # ✅ TEXTE TRADUIT FRANÇAIS (LTR)
            doc.add_heading('🌍 Texte Traduit (Français)', level=1)
            WordExporter._add_formatted_text(doc, translated_text, is_arabic=False)
            
            # Section statistiques
            doc.add_page_break()
            stats_heading = doc.add_heading('📊 Statistiques de Traduction', level=1)
            stats_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            stats_para = doc.add_paragraph()
            stats_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats_text = (f"• Caractères originaux: {len(original_text)}\n"
                         f"• Caractères traduits: {len(translated_text)}\n"
                         f"• Ratio de compression: {len(translated_text)/max(1, len(original_text)):.2f}\n"
                         f"• Date de traduction: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
                         f"• Modèle: DeepSeek AI")
            stats_para.add_run(stats_text)
            
            # Pied de page
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("Traduction exportée par Plateforme Éducative Arabe V2")
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            footer_run.italic = True
            
            # Sauvegarde
            filepath = os.path.join(Config.UPLOAD_FOLDER, filename)
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            doc.save(filepath)
            logger.info(f"✅ Export traduction Word réussi: {filepath}")
            logger.info(f"📁 Taille du fichier: {os.path.getsize(filepath)} bytes")
            
            return filepath
            
        except Exception as e:
            logger.error(f"❌ Erreur export traduction Word: {e}", exc_info=True)
            raise

    @staticmethod
    def _add_formatted_text(doc, text, is_arabic=False):
        """Ajoute du texte formaté avec sauts de ligne préservés"""
        lines = text.split('\n')
        current_page = None
        page_count = 0
        
        logger.info(f"📝 Ajout de {len(lines)} lignes (arabe: {is_arabic})")
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                doc.add_paragraph()  # Ligne vide
                continue
                
            # Détection des en-têtes de page
            if line.startswith('--- Page') and '---' in line:
                page_count += 1
                current_page = line.replace('---', '').strip()
                
                # Saut de page (sauf pour la première page)
                if page_count > 1:
                    doc.add_page_break()
                
                # En-tête de page stylisé
                page_para = doc.add_paragraph()
                page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                page_para.paragraph_format.space_before = Pt(12)
                page_para.paragraph_format.space_after = Pt(6)
                
                page_run = page_para.add_run(f"📄 {current_page}")
                page_run.bold = True
                page_run.font.size = Pt(12)
                page_color = RGBColor(0, 0, 128) if is_arabic else RGBColor(0, 100, 0)
                page_run.font.color.rgb = page_color
                
            else:
                # Texte normal
                paragraph = doc.add_paragraph()
                
                if is_arabic:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    WordExporter._set_paragraph_rtl(paragraph)
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.space_before = Pt(0)
                
                run = paragraph.add_run(line)
                
                if is_arabic:
                    run.font.name = 'Traditional Arabic'
                    run.font.size = Pt(11)
                else:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    # Style pour le français
                    if len(line) < 100 and line.endswith(':'):
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 100, 0)

    @staticmethod
    def _set_paragraph_rtl(paragraph):
        """Configure un paragraphe en mode Right-to-Left (pour l'arabe)"""
        try:
            # Configuration RTL via les propriétés OOXML
            p_pr = paragraph._p.get_or_add_pPr()
            p_pr.append(OxmlElement('w:bidi'))
            # Définir la direction RTL
            p_pr.append(OxmlElement('w:rtl'))
        except Exception as e:
            logger.warning(f"⚠️ Impossible de configurer RTL: {e}")
            # Fallback: simple alignement à droite
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    @staticmethod
    def test_export():
        """Méthode de test pour vérifier l'export"""
        try:
            test_ocr = """--- Page 1 ---
هذا هو السطر الأول
هذا هو السطر الثاني

هذا سطر بعد مسافة

--- Page 2 ---
سطر من صفحة جديدة
آخر سطر"""

            test_translation = """--- Page 1 ---
This is the first line
This is the second line

This is a line after space

--- Page 2 ---
Line from new page
Last line"""

            context = {
                'auteur': 'Test Author',
                'titre': 'Test Document',
                'sujet': 'Test Subject',
                'genre': 'Littérature',
                'niveau_langue': 'Standard'
            }

            WordExporter.export_ocr_to_word(test_ocr, "test_ocr.docx")
            WordExporter.export_translation_to_word(test_ocr, test_translation, context, "test_trad.docx")
            
            print("✅ Tests d'export Word réussis!")
            return True
            
        except Exception as e:
            print(f"❌ Erreur test export: {e}")
            return False

if __name__ == "__main__":
    WordExporter.test_export()